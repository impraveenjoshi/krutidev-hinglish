"""
Word-style editor shell: ribbon UI, page layout, header/footer, images,
plus Hinglish → Kruti Dev encoding.
"""

from __future__ import annotations

import sys
from pathlib import Path

from PyQt6.QtCore import QMarginsF, QRectF, QSizeF, Qt, QTimer
from PyQt6.QtGui import (
    QAction,
    QColor,
    QFont,
    QFontDatabase,
    QGuiApplication,
    QKeySequence,
    QPageLayout,
    QPageSize,
    QPainter,
    QPalette,
    QTextBlockFormat,
    QTextCharFormat,
    QTextCursor,
    QTextDocument,
    QTextImageFormat,
)
from PyQt6.QtPrintSupport import QPrintDialog, QPrinter, QPrintPreviewDialog
from PyQt6.QtWidgets import (
    QApplication,
    QCheckBox,
    QComboBox,
    QFileDialog,
    QFontComboBox,
    QFrame,
    QGraphicsDropShadowEffect,
    QHBoxLayout,
    QLabel,
    QMainWindow,
    QMenuBar,
    QMessageBox,
    QPushButton,
    QScrollArea,
    QSizePolicy,
    QSpinBox,
    QTabWidget,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

try:
    from docx import Document
except ImportError:
    Document = None

from hinglish_kruti import hinglish_document_to_krutidev, hinglish_selection_to_krutidev

# Paper presets: label -> QPageSize.PageSizeId
PAPER_PRESETS: list[tuple[str, QPageSize.PageSizeId]] = [
    ("A4", QPageSize.PageSizeId.A4),
    ("A3", QPageSize.PageSizeId.A3),
    ("A5", QPageSize.PageSizeId.A5),
    ("B4 (JIS)", QPageSize.PageSizeId.B4),
    ("B5 (JIS)", QPageSize.PageSizeId.B5),
    ("Letter", QPageSize.PageSizeId.Letter),
    ("Legal", QPageSize.PageSizeId.Legal),
    ("Tabloid / Ledger", QPageSize.PageSizeId.Tabloid),
    ("Executive", QPageSize.PageSizeId.Executive),
]


def resolve_kruti_family() -> str:
    """Prefer a bundled TTF, then installed Windows Kruti Dev faces (names vary)."""

    base = Path(__file__).resolve().parent
    for rel in ("fonts/KrutiDev010.ttf", "fonts/Kruti_Dev_010.ttf", "KrutiDev010.ttf"):
        p = base / rel
        if p.is_file():
            fid = QFontDatabase.addApplicationFont(str(p))
            if fid >= 0:
                fams = QFontDatabase.applicationFontFamilies(fid)
                if fams:
                    return fams[0]

    installed = list(QFontDatabase.families())
    by_lower = {f.lower(): f for f in installed}
    for name in (
        "Kruti Dev",
        "Kruti Dev 010",
        "KrutiDev",
        "KrutiDev010",
        "KD",
        "Krutidev",
    ):
        if name in installed:
            return name
        hit = by_lower.get(name.lower())
        if hit:
            return hit
    for f in installed:
        if "kruti" in f.lower():
            return f
    return ""


def cm_to_px(cm: float, dpi: float) -> float:
    return (cm / 2.54) * dpi


# Keeps ribbon/menus readable when Windows is in dark mode (avoids white text on light QSS).
APP_CHROME_STYLESHEET = """
QMainWindow { background: #f0f0f0; color: #1a1a1a; }
QMenuBar {
    background: #f3f3f3;
    color: #1a1a1a;
    border-bottom: 1px solid #c0c0c0;
    padding: 2px;
}
QMenuBar::item { padding: 4px 12px; background: transparent; color: #1a1a1a; }
QMenuBar::item:selected { background: #ddeeff; color: #1a1a1a; }
QMenu {
    background: #ffffff;
    color: #1a1a1a;
    border: 1px solid #a6a6a6;
}
QMenu::item:selected { background: #0078d4; color: #ffffff; }

QToolBar {
    background: #f3f3f3;
    color: #1a1a1a;
    border: none;
    border-bottom: 1px solid #c0c0c0;
    spacing: 6px;
    padding: 6px 4px;
}
QToolBar QLabel { color: #303030; font-size: 9pt; }

QToolButton {
    background: #ffffff;
    color: #1a1a1a;
    border: 1px solid #b8b8b8;
    border-radius: 3px;
    padding: 6px 10px;
    min-width: 26px;
    min-height: 24px;
    font-weight: 600;
}
QToolButton:hover { background: #e5f3ff; border-color: #0078d4; color: #000000; }
QToolButton:pressed { background: #cce8ff; }
QToolButton:disabled { color: #888888; background: #f0f0f0; }

QTabWidget::pane { border-top: 1px solid #c0c0c0; background: #f3f3f3; top: -1px; }
QTabBar::tab {
    background: #d4d4d4;
    color: #1a1a1a;
    padding: 8px 18px;
    margin-right: 2px;
    border: 1px solid #9a9a9a;
    border-bottom: none;
    border-top-left-radius: 4px;
    border-top-right-radius: 4px;
    font-weight: 500;
}
QTabBar::tab:selected {
    background: #f3f3f3;
    color: #000000;
    border-bottom: 2px solid #0078d4;
    font-weight: 700;
}
QTabBar::tab:hover:!selected { background: #e8e8e8; color: #000000; }

QComboBox, QFontComboBox {
    background: #ffffff;
    color: #1a1a1a;
    border: 1px solid #8a8a8a;
    border-radius: 3px;
    padding: 4px 8px;
    min-height: 24px;
    selection-background-color: #0078d4;
    selection-color: #ffffff;
}
QComboBox QAbstractItemView, QFontComboBox QAbstractItemView {
    background: #ffffff;
    color: #1a1a1a;
    selection-background-color: #0078d4;
    selection-color: #ffffff;
}
QComboBox::drop-down { border: none; width: 22px; }
QFontComboBox::drop-down { border: none; width: 22px; }

QLabel { color: #303030; }

QPushButton {
    background: #ffffff;
    color: #1a1a1a;
    border: 1px solid #8a8a8a;
    border-radius: 3px;
    padding: 6px 14px;
    min-height: 24px;
    font-weight: 600;
}
QPushButton:hover { background: #e5f3ff; border-color: #0078d4; }
QPushButton:pressed { background: #cce8ff; }

QCheckBox { color: #1a1a1a; spacing: 8px; font-size: 9pt; }
QCheckBox::indicator { width: 16px; height: 16px; }

QSpinBox {
    background: #ffffff;
    color: #1a1a1a;
    border: 1px solid #8a8a8a;
    border-radius: 3px;
    padding: 4px;
    min-height: 24px;
}

QStatusBar {
    background: #dedede;
    color: #1a1a1a;
    border-top: 1px solid #a0a0a0;
    padding: 5px 8px;
    font-size: 10pt;
}
QStatusBar::item { border: none; }

#RibbonGroupChunk {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 #ffffff, stop:1 #eeeeee);
    border: 1px solid #b0b0b0;
    border-radius: 6px;
}
#PageSheet {
    background: #ffffff;
    border: 1px solid #9aa0a6;
    border-radius: 3px;
}
QTextEdit {
    color: #1a1a1a;
}
QScrollBar:vertical { background: #e8e8e8; width: 12px; margin: 0; }
QScrollBar::handle:vertical { background: #b0b0b0; min-height: 24px; border-radius: 4px; }
QScrollBar:horizontal { background: #e8e8e8; height: 12px; margin: 0; }
QScrollBar::handle:horizontal { background: #b0b0b0; min-width: 24px; border-radius: 4px; }
"""


def apply_light_chrome(app: QApplication) -> None:
    pal = QPalette()
    c = QColor
    cr = QPalette.ColorRole
    pal.setColor(cr.Window, c(240, 240, 240))
    pal.setColor(cr.WindowText, c(26, 26, 26))
    pal.setColor(cr.Base, c(255, 255, 255))
    pal.setColor(cr.AlternateBase, c(248, 248, 248))
    pal.setColor(cr.Text, c(26, 26, 26))
    pal.setColor(cr.Button, c(237, 237, 237))
    pal.setColor(cr.ButtonText, c(26, 26, 26))
    pal.setColor(cr.ToolTipBase, c(255, 255, 220))
    pal.setColor(cr.ToolTipText, c(26, 26, 26))
    pal.setColor(cr.Highlight, c(0, 120, 212))
    pal.setColor(cr.HighlightedText, c(255, 255, 255))
    pal.setColor(cr.Mid, c(184, 184, 184))
    pal.setColor(cr.Dark, c(106, 106, 106))
    app.setPalette(pal)
    app.setStyleSheet(APP_CHROME_STYLESHEET)


class RibbonGroup(QFrame):
    """One labeled group in the ribbon (Office-style chunk)."""

    def __init__(self, title: str, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.setObjectName("RibbonGroupChunk")
        self.setFrameShape(QFrame.Shape.StyledPanel)
        outer = QVBoxLayout(self)
        outer.setContentsMargins(8, 6, 8, 4)
        outer.setSpacing(4)
        self._row = QHBoxLayout()
        self._row.setSpacing(6)
        outer.addLayout(self._row)
        lab = QLabel(title)
        lab.setStyleSheet("color: #444444; font-size: 9pt; font-weight: 600;")
        outer.addWidget(lab, alignment=Qt.AlignmentFlag.AlignHCenter)

    def add_widget(self, w: QWidget) -> None:
        self._row.addWidget(w)


class MainWindow(QMainWindow):
    def __init__(self) -> None:
        super().__init__()
        self.setWindowTitle("Document — Hinglish → Kruti Dev")
        self.resize(1100, 780)

        self.kruti_family = resolve_kruti_family()
        self._warned_missing_kruti_font = False

        self._paper_id: QPageSize.PageSizeId = QPageSize.PageSizeId.A4
        self._landscape = False
        self._margins_cm = {"L": 2.0, "R": 2.0, "T": 2.5, "B": 2.5}
        self._footer_refresh_busy = False
        self._footer_template_plain = ""

        self._build_menu_bar()
        self._build_central_page_editor()
        self._build_ribbon()
        self._build_status_hint()

        self._apply_document_page_geometry()
        self._sync_font_combo_with_editor()

        self.editor.textChanged.connect(self._schedule_footer_refresh)
        self.editor.cursorPositionChanged.connect(self._schedule_footer_refresh)
        self.editor.textChanged.connect(self._update_page_indicator)
        self.editor.cursorPositionChanged.connect(self._update_page_indicator)

        self._footer_timer = QTimer(self)
        self._footer_timer.setSingleShot(True)
        self._footer_timer.timeout.connect(self._refresh_footer_placeholders)

    # --- Menu bar ---

    def _build_menu_bar(self) -> None:
        mb = QMenuBar(self)
        self.setMenuBar(mb)
        file_menu = mb.addMenu("File")
        a_open = QAction("Open…", self)
        a_open.setShortcut(QKeySequence.StandardKey.Open)
        a_open.triggered.connect(self._open_file)
        file_menu.addAction(a_open)
        a_save = QAction("Save…", self)
        a_save.setShortcut(QKeySequence.StandardKey.Save)
        a_save.triggered.connect(self._save_file)
        file_menu.addAction(a_save)
        file_menu.addSeparator()
        a_pv = QAction("Print preview…", self)
        a_pv.triggered.connect(self._print_preview)
        file_menu.addAction(a_pv)
        a_pr = QAction("Print…", self)
        a_pr.setShortcut(QKeySequence.StandardKey.Print)
        a_pr.triggered.connect(self._print_dialog)
        file_menu.addAction(a_pr)
        file_menu.addSeparator()
        a_exit = QAction("Exit", self)
        a_exit.setShortcut(QKeySequence.StandardKey.Quit)
        a_exit.triggered.connect(self.close)
        file_menu.addAction(a_exit)

    # --- Ribbon ---

    def _build_ribbon(self) -> None:
        tabs = QTabWidget()
        tabs.setDocumentMode(True)
        tabs.setMovable(False)

        home = QWidget()
        hlay = QHBoxLayout(home)
        hlay.setContentsMargins(10, 8, 10, 8)
        hlay.setSpacing(10)

        clip = RibbonGroup("Clipboard")
        paste_btn = QPushButton("Paste")
        paste_btn.setToolTip("Paste (Ctrl+V)")
        paste_btn.clicked.connect(self.editor.paste)
        clip.add_widget(paste_btn)
        save_btn = QPushButton("Save")
        save_btn.setToolTip("Save (Ctrl+S)")
        save_btn.clicked.connect(self._save_file)
        clip.add_widget(save_btn)
        hlay.addWidget(clip)

        font_g = RibbonGroup("Font")
        self.font_combo = QFontComboBox()
        self.font_combo.setMaximumWidth(220)
        self.font_combo.currentFontChanged.connect(self._apply_font_family)
        font_g.add_widget(self.font_combo)
        self.size_combo = QComboBox()
        for s in range(8, 73, 1):
            self.size_combo.addItem(str(s), s)
        self.size_combo.setCurrentText("14")
        self.size_combo.setMaximumWidth(56)
        self.size_combo.currentIndexChanged.connect(self._apply_size)
        font_g.add_widget(QLabel("Size"))
        font_g.add_widget(self.size_combo)
        for label, slot, key in (
            ("B", self._toggle_bold, "Ctrl+B"),
            ("I", self._toggle_italic, "Ctrl+I"),
            ("U", self._toggle_underline, "Ctrl+U"),
        ):
            act = QAction(label, self)
            act.setShortcut(key)
            act.triggered.connect(slot)
            tb = QFrame()
            tlay = QVBoxLayout(tb)
            tlay.addWidget(self._tool_button_from_action(act))
            font_g.add_widget(tb)
        hlay.addWidget(font_g)

        para = RibbonGroup("Paragraph")
        for text, align in (
            ("L", Qt.AlignmentFlag.AlignLeft),
            ("C", Qt.AlignmentFlag.AlignHCenter),
            ("R", Qt.AlignmentFlag.AlignRight),
            ("J", Qt.AlignmentFlag.AlignJustify),
        ):
            act = QAction(text, self)
            act.triggered.connect(lambda _=False, a=align: self._set_alignment(a))
            para.add_widget(self._tool_button_from_action(act))
        hlay.addWidget(para)

        kr = RibbonGroup("Kruti / Hindi")
        a1 = QAction("Selection → Kruti", self)
        a1.setShortcut("Ctrl+Shift+H")
        a1.triggered.connect(self._convert_selection)
        a2 = QAction("Latin words → Kruti", self)
        a2.triggered.connect(self._convert_document)
        kr.add_widget(self._tool_button_from_action(a1))
        kr.add_widget(self._tool_button_from_action(a2))
        hlay.addWidget(kr)
        hlay.addStretch(1)

        insert = QWidget()
        ilay = QHBoxLayout(insert)
        ilay.setContentsMargins(10, 8, 10, 8)
        ilay.setSpacing(10)
        ins = RibbonGroup("Insert")
        pic = QAction("Picture", self)
        pic.triggered.connect(self._insert_picture)
        brk = QAction("Page break", self)
        brk.triggered.connect(self._insert_page_break)
        ins.add_widget(self._tool_button_from_action(pic))
        ins.add_widget(self._tool_button_from_action(brk))
        ilay.addWidget(ins)
        ilay.addStretch(1)

        layout_w = QWidget()
        llay = QHBoxLayout(layout_w)
        llay.setContentsMargins(10, 8, 10, 8)
        llay.setSpacing(10)

        page_g = RibbonGroup("Page")
        self.paper_combo = QComboBox()
        for label, pid in PAPER_PRESETS:
            self.paper_combo.addItem(label, pid)
        self.paper_combo.setCurrentIndex(0)
        self.paper_combo.currentIndexChanged.connect(self._on_paper_changed)
        page_g.add_widget(QLabel("Size"))
        page_g.add_widget(self.paper_combo)
        self.orient_combo = QComboBox()
        self.orient_combo.addItem("Portrait", False)
        self.orient_combo.addItem("Landscape", True)
        self.orient_combo.currentIndexChanged.connect(self._on_orient_changed)
        page_g.add_widget(QLabel("Orientation"))
        page_g.add_widget(self.orient_combo)
        llay.addWidget(page_g)

        margin_g = RibbonGroup("Margins (cm)")
        self.margin_spin: dict[str, QSpinBox] = {}
        for key, lab in (("L", "L"), ("R", "R"), ("T", "T"), ("B", "B")):
            margin_g.add_widget(QLabel(lab))
            sp = QSpinBox()
            sp.setRange(0, 99)
            sp.setValue(int(self._margins_cm[key]))
            sp.setSuffix(" cm")
            sp.valueChanged.connect(self._on_margins_changed)
            self.margin_spin[key] = sp
            margin_g.add_widget(sp)
        llay.addWidget(margin_g)

        hf_g = RibbonGroup("Header / footer")
        self.show_header_cb = QCheckBox("Header")
        self.show_footer_cb = QCheckBox("Footer")
        self.show_header_cb.toggled.connect(self._on_header_footer_toggled)
        self.show_footer_cb.toggled.connect(self._on_header_footer_toggled)
        hf_g.add_widget(self.show_header_cb)
        hf_g.add_widget(self.show_footer_cb)
        hf_g.add_widget(QLabel("Placeholders:"))
        hf_g.add_widget(QLabel("[[page]] [[pages]]"))
        llay.addWidget(hf_g)

        pr_g = RibbonGroup("Print")
        apv = QAction("Preview", self)
        apv.triggered.connect(self._print_preview)
        apr = QAction("Print", self)
        apr.triggered.connect(self._print_dialog)
        pr_g.add_widget(self._tool_button_from_action(apv))
        pr_g.add_widget(self._tool_button_from_action(apr))
        llay.addWidget(pr_g)
        llay.addStretch(1)

        tabs.addTab(home, "Home")
        tabs.addTab(insert, "Insert")
        tabs.addTab(layout_w, "Layout")

        tb = self.addToolBar("ribbon")
        tb.setMovable(False)
        tb.setFloatable(False)
        tb.setIconSize(tb.iconSize())
        w = QWidget()
        wl = QVBoxLayout(w)
        wl.setContentsMargins(0, 0, 0, 0)
        wl.addWidget(tabs)
        tb.addWidget(w)

    def _tool_button_from_action(self, act: QAction) -> QWidget:
        from PyQt6.QtWidgets import QToolButton

        b = QToolButton()
        b.setDefaultAction(act)
        b.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextOnly)
        b.setAutoRaise(True)
        return b

    # --- Central: header + page + footer ---

    def _build_central_page_editor(self) -> None:
        outer = QWidget()
        ov = QVBoxLayout(outer)
        ov.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        scroll.setStyleSheet("QScrollArea { background: #c8c8c8; border: none; }")

        canvas = QWidget()
        canvas.setStyleSheet("background: #c8c8c8;")
        ch = QHBoxLayout(canvas)
        ch.setContentsMargins(24, 24, 24, 24)
        ch.addStretch(1)

        col = QVBoxLayout()
        col.setSpacing(14)
        col.setAlignment(Qt.AlignmentFlag.AlignTop)

        self.header_edit = QTextEdit()
        self.header_edit.setAcceptRichText(True)
        self.header_edit.setPlaceholderText("Header (optional). Insert → Picture works here when header has focus.")
        self.header_edit.setMaximumHeight(72)
        self.header_edit.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.header_edit.setStyleSheet(
            "QTextEdit { background: #fff; border: 1px solid #adadad; font-size: 10pt; }"
        )
        self.header_edit.hide()

        self.editor = QTextEdit()
        self.editor.setObjectName("body")
        self.editor.setAcceptRichText(True)
        self.editor.setPlaceholderText("Document text…")
        self.editor.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
        self.editor.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.editor.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.editor.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Expanding)
        self.editor.setStyleSheet(
            "QTextEdit#body { background: #ffffff; border: none; font-size: 11pt; "
            "padding: 8px; }"
        )

        page_sheet = QFrame()
        page_sheet.setObjectName("PageSheet")
        ps_lay = QVBoxLayout(page_sheet)
        ps_lay.setContentsMargins(0, 0, 0, 0)
        ps_lay.addWidget(self.editor)
        shadow = QGraphicsDropShadowEffect(page_sheet)
        shadow.setBlurRadius(28)
        shadow.setColor(QColor(0, 0, 0, 55))
        shadow.setOffset(2, 3)
        page_sheet.setGraphicsEffect(shadow)

        self.footer_edit = QTextEdit()
        self.footer_edit.setAcceptRichText(True)
        self.footer_edit.setPlaceholderText(
            "Footer (optional). For live numbers use plain text with [[page]] and [[pages]] "
            "(images in footer are not supported together with those tokens)."
        )
        self.footer_edit.setMaximumHeight(72)
        self.footer_edit.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.footer_edit.setStyleSheet(
            "QTextEdit { background: #fff; border: 1px solid #adadad; font-size: 10pt; }"
        )
        self.footer_edit.hide()
        self.footer_edit.textChanged.connect(self._on_footer_user_edit)

        col.addWidget(self.header_edit)
        col.addWidget(page_sheet, 1)
        col.addWidget(self.footer_edit)

        ch.addLayout(col)
        ch.addStretch(1)
        scroll.setWidget(canvas)
        ov.addWidget(scroll, 1)
        self.setCentralWidget(outer)

        if self.kruti_family:
            self.editor.setFont(QFont(self.kruti_family, 14))
        else:
            self.editor.setFont(QFont("Segoe UI", 11))

    def _build_status_hint(self) -> None:
        if self.kruti_family:
            msg = f'Kruti: "{self.kruti_family}" — add fonts/KrutiDev010.ttf beside the app if missing.'
        else:
            msg = "Kruti Dev not found — converted text needs Kruti font to display as Hindi."
        self.statusBar().showMessage(msg, 6000)

        self._status_page_label = QLabel("Page —")
        self._status_page_label.setStyleSheet(
            "color: #111; font-weight: 600; padding: 2px 12px; background: #d0d0d0; "
            "border-radius: 3px; margin: 2px;"
        )
        self.statusBar().addPermanentWidget(self._status_page_label)
        self._update_page_indicator()

    # --- Geometry / page ---

    def _dpi(self) -> float:
        scr = QGuiApplication.primaryScreen()
        return float(scr.logicalDotsPerInchX()) if scr else 96.0

    def _page_size_points(self) -> tuple[float, float]:
        ps = QPageSize(self._paper_id)
        sz = ps.sizePoints()
        w, h = float(sz.width()), float(sz.height())
        if self._landscape:
            return h, w
        return w, h

    def _apply_document_page_geometry(self) -> None:
        w_pt, h_pt = self._page_size_points()
        doc = self.editor.document()
        doc.setPageSize(QSizeF(w_pt, h_pt))

        dpi = self._dpi()
        page_px = int(round((w_pt / 72.0) * dpi))
        self.editor.setFixedWidth(max(200, page_px))

        self._apply_frame_margins()
        self._schedule_footer_refresh()
        self._update_page_indicator()

    def _apply_frame_margins(self) -> None:
        doc = self.editor.document()
        root = doc.rootFrame()
        fmt = root.frameFormat()
        dpi = self._dpi()
        fmt.setLeftMargin(cm_to_px(self._margins_cm["L"], dpi))
        fmt.setRightMargin(cm_to_px(self._margins_cm["R"], dpi))
        fmt.setTopMargin(cm_to_px(self._margins_cm["T"], dpi))
        fmt.setBottomMargin(cm_to_px(self._margins_cm["B"], dpi))
        root.setFrameFormat(fmt)

    def _on_paper_changed(self) -> None:
        pid = self.paper_combo.currentData()
        if pid is not None:
            self._paper_id = pid
        self._apply_document_page_geometry()

    def _on_orient_changed(self) -> None:
        self._landscape = bool(self.orient_combo.currentData())
        self._apply_document_page_geometry()

    def _on_margins_changed(self) -> None:
        self._margins_cm["L"] = float(self.margin_spin["L"].value())
        self._margins_cm["R"] = float(self.margin_spin["R"].value())
        self._margins_cm["T"] = float(self.margin_spin["T"].value())
        self._margins_cm["B"] = float(self.margin_spin["B"].value())
        self._apply_frame_margins()

    def _on_header_footer_toggled(self) -> None:
        self.header_edit.setVisible(self.show_header_cb.isChecked())
        self.footer_edit.setVisible(self.show_footer_cb.isChecked())
        if self.show_footer_cb.isChecked() and not self._footer_template_plain.strip():
            self._footer_template_plain = "Page [[page]] of [[pages]]"
        if self.show_footer_cb.isChecked():
            self._refresh_footer_placeholders()

    # --- Footer placeholders ---

    def _on_footer_user_edit(self) -> None:
        if self._footer_refresh_busy:
            return
        self._footer_template_plain = self.footer_edit.toPlainText()
        self._schedule_footer_refresh()

    def _schedule_footer_refresh(self) -> None:
        if not self.show_footer_cb.isChecked():
            return
        self._footer_timer.start(120)

    def _refresh_footer_placeholders(self) -> None:
        if not self.show_footer_cb.isChecked():
            return
        doc = self.editor.document()
        pages = max(1, doc.pageCount())
        page = max(1, min(pages, self._current_page_number()))

        tmpl = self._footer_template_plain
        if "[[page]]" not in tmpl and "[[pages]]" not in tmpl:
            return
        new_plain = tmpl.replace("[[pages]]", str(pages)).replace("[[page]]", str(page))
        if self.footer_edit.toPlainText() == new_plain:
            return
        self._footer_refresh_busy = True
        self.footer_edit.blockSignals(True)
        try:
            self.footer_edit.setPlainText(new_plain)
        finally:
            self.footer_edit.blockSignals(False)
            self._footer_refresh_busy = False

    def _document_y_for_position(self, position: int) -> float:
        """Precise document Y for position using layout hit-test (for accurate page calc)."""
        doc = self.editor.document()
        layout = doc.documentLayout()
        if layout is None:
            return 0.0
        block = doc.findBlock(position)
        if not block.isValid():
            return 0.0
        br = layout.blockBoundingRect(block)
        lay = block.layout()
        if lay is None or not block.isVisible():
            return max(0.0, br.top())
        rel = position - block.position()
        rel = max(0, min(rel, max(0, block.length() - 1)))
        line = lay.lineForTextPosition(rel)
        if line.isValid():
            return max(0.0, br.top() + line.y())
        return max(0.0, br.top())

    def _current_page_number(self) -> int:
        """Get current page number for cursor position (1-based)."""
        doc = self.editor.document()
        ph = float(doc.pageSize().height())
        if ph <= 0:
            return 1
        pos = self.editor.textCursor().position()
        y = self._document_y_for_position(pos)
        page_num = max(1, int(y // ph) + 1)
        total_pages = max(1, doc.pageCount())
        return min(page_num, total_pages)

    def _update_page_indicator(self) -> None:
        if not hasattr(self, "_status_page_label"):
            return
        doc = self.editor.document()
        n = max(1, doc.pageCount())
        cur = self._current_page_number()
        self._status_page_label.setText(f"Page {cur} / {n}")

    # --- Print ---

    def _clone_body_document(self) -> QTextDocument:
        d = QTextDocument()
        d.setHtml(self.editor.toHtml())
        d.setDefaultFont(self.editor.document().defaultFont())
        w_pt, h_pt = self._page_size_points()
        d.setPageSize(QSizeF(w_pt, h_pt))
        d.rootFrame().setFrameFormat(self.editor.document().rootFrame().frameFormat())
        return d

    def _configured_printer(self) -> QPrinter:
        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        layout = printer.pageLayout()
        layout.setPageSize(QPageSize(self._paper_id))
        layout.setOrientation(
            QPageLayout.Orientation.Landscape
            if self._landscape
            else QPageLayout.Orientation.Portrait
        )
        m = self._margins_cm
        layout.setMargins(QMarginsF(m["L"] * 10.0, m["T"] * 10.0, m["R"] * 10.0, m["B"] * 10.0))
        layout.setUnits(QPageLayout.Unit.Millimeter)
        printer.setPageLayout(layout)
        return printer

    def _render_document_to_printer(self, printer: QPrinter) -> None:
        print_doc = self._clone_body_document()
        pages = max(1, print_doc.pageCount())
        pw = float(print_doc.pageSize().width())
        ph = float(print_doc.pageSize().height())

        painter = QPainter()
        if not painter.begin(printer):
            QMessageBox.warning(self, "Print", "Could not start printing.")
            return
        painter.setRenderHint(QPainter.RenderHint.TextAntialiasing)

        page_px = printer.pageRect(QPrinter.Unit.DevicePixel)
        dpi_x = max(1, printer.logicalDpiX())
        dpi_y = max(1, printer.logicalDpiY())

        def pt_to_px_x(pt: float) -> float:
            return pt * dpi_x / 72.0

        def pt_to_px_y(pt: float) -> float:
            return pt * dpi_y / 72.0

        header_h_px = pt_to_px_y(30.0) if self.show_header_cb.isChecked() else 0.0
        footer_h_px = pt_to_px_y(26.0) if self.show_footer_cb.isChecked() else 0.0

        inner = QRectF(page_px)
        body = QRectF(
            inner.x(),
            inner.y() + header_h_px,
            inner.width(),
            max(1.0, inner.height() - header_h_px - footer_h_px),
        )

        pw_px = pt_to_px_x(pw)
        scale = body.width() / max(1.0, pw_px)

        header_text = self.header_edit.toPlainText() if self.show_header_cb.isChecked() else ""
        footer_tmpl = (self._footer_template_plain or "").strip()
        if self.show_footer_cb.isChecked() and not footer_tmpl:
            footer_tmpl = self.footer_edit.toPlainText()

        head_font = QFont("Segoe UI", 9)
        foot_font = QFont("Segoe UI", 9)

        for i in range(pages):
            if i > 0:
                printer.newPage()

            if header_h_px > 0 and header_text.strip():
                painter.setFont(head_font)
                r_h = QRectF(inner.x(), inner.y(), inner.width(), header_h_px)
                painter.drawText(
                    r_h,
                    int(Qt.AlignmentFlag.AlignHCenter | Qt.AlignmentFlag.AlignVCenter),
                    header_text,
                )

            if footer_h_px > 0 and footer_tmpl.strip():
                painter.setFont(foot_font)
                ftxt = footer_tmpl.replace("[[pages]]", str(pages)).replace("[[page]]", str(i + 1))
                r_f = QRectF(inner.x(), inner.bottom() - footer_h_px, inner.width(), footer_h_px)
                painter.drawText(
                    r_f,
                    int(Qt.AlignmentFlag.AlignHCenter | Qt.AlignmentFlag.AlignVCenter),
                    ftxt,
                )

            painter.save()
            painter.translate(body.x(), body.y())
            painter.scale(scale, scale)
            painter.translate(0.0, -float(i) * ph)
            print_doc.drawContents(painter, QRectF(0.0, float(i) * ph, pw, ph))
            painter.restore()

        painter.end()

    def _print_preview(self) -> None:
        printer = self._configured_printer()
        dlg = QPrintPreviewDialog(printer, self)
        dlg.setWindowTitle("Print preview")
        dlg.paintRequested.connect(self._render_document_to_printer)
        dlg.exec()

    def _print_dialog(self) -> None:
        printer = self._configured_printer()
        dlg = QPrintDialog(printer, self)
        if dlg.exec() != QPrintDialog.DialogCode.Accepted:
            return
        self._render_document_to_printer(printer)

    # --- Formatting ---

    def _current_point_size(self) -> int:
        return int(self.size_combo.currentData()) if self.size_combo.currentData() else 14

    def _sync_font_combo_with_editor(self) -> None:
        self.font_combo.blockSignals(True)
        try:
            if self.kruti_family:
                self.font_combo.setCurrentFont(QFont(self.kruti_family, self._current_point_size()))
            else:
                self.font_combo.setCurrentFont(self.editor.font())
        finally:
            self.font_combo.blockSignals(False)

    def _apply_kruti_display_after_convert(self) -> None:
        if self.kruti_family:
            f = QFont(self.kruti_family, self._current_point_size())
            self.editor.setFont(f)
            self.font_combo.blockSignals(True)
            try:
                self.font_combo.setCurrentFont(f)
            finally:
                self.font_combo.blockSignals(False)
            return
        if self._warned_missing_kruti_font:
            return
        self._warned_missing_kruti_font = True
        QMessageBox.information(
            self,
            "Kruti Dev font needed to view Hindi",
            "The conversion produced Kruti Dev *encoded* text. Use the Kruti Dev font "
            "(or fonts/KrutiDev010.ttf beside the .exe) to see Devanagari shapes.",
        )

    def _merge_format(self, fmt: QTextCharFormat) -> None:
        c = self.editor.textCursor()
        if not c.hasSelection():
            c.select(QTextCursor.SelectionType.WordUnderCursor)
        c.mergeCharFormat(fmt)
        self.editor.mergeCurrentCharFormat(fmt)

    def _toggle_bold(self) -> None:
        fmt = QTextCharFormat()
        w = QFont.Weight.Bold if self.editor.fontWeight() != QFont.Weight.Bold else QFont.Weight.Normal
        fmt.setFontWeight(w)
        self._merge_format(fmt)

    def _toggle_italic(self) -> None:
        fmt = QTextCharFormat()
        fmt.setFontItalic(not self.editor.fontItalic())
        self._merge_format(fmt)

    def _toggle_underline(self) -> None:
        fmt = QTextCharFormat()
        fmt.setFontUnderline(not self.editor.fontUnderline())
        self._merge_format(fmt)

    def _apply_size(self) -> None:
        fmt = QTextCharFormat()
        fmt.setFontPointSize(float(self._current_point_size()))
        self._merge_format(fmt)

    def _apply_font_family(self, font: QFont) -> None:
        fmt = QTextCharFormat()
        fmt.setFontFamily(font.family())
        self._merge_format(fmt)

    def _set_alignment(self, align: Qt.AlignmentFlag) -> None:
        c = self.editor.textCursor()
        bf = QTextBlockFormat()
        bf.setAlignment(align)
        c.mergeBlockFormat(bf)

    def _insert_picture(self) -> None:
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Insert picture",
            "",
            "Images (*.png *.jpg *.jpeg *.bmp *.gif *.webp);;All (*.*)",
        )
        if not path:
            return
        target = self.editor
        if self.header_edit.hasFocus():
            target = self.header_edit
        elif self.footer_edit.hasFocus():
            target = self.footer_edit

        c = target.textCursor()
        fmt = QTextImageFormat()
        fmt.setName(path)
        fmt.setWidth(320)
        c.insertImage(fmt)
        target.setTextCursor(c)
        if target is self.editor:
            self.editor.setFocus()

    def _insert_page_break(self) -> None:
        c = self.editor.textCursor()
        bf = QTextBlockFormat()
        bf.setPageBreakPolicy(QTextBlockFormat.PageBreakFlag.PageBreak_AlwaysAfter)
        c.mergeBlockFormat(bf)
        c.insertBlock()
        self.editor.setTextCursor(c)

    def _kruti_char_format(self) -> QTextCharFormat:
        fmt = QTextCharFormat()
        if self.kruti_family:
            fmt.setFontFamily(self.kruti_family)
        fmt.setFontPointSize(float(self._current_point_size()))
        return fmt

    def _convert_selection(self) -> None:
        c = self.editor.textCursor()
        if not c.hasSelection():
            QMessageBox.information(self, "Nothing selected", "Select Hinglish text first.")
            return
        raw = c.selectedText().replace("\u2029", "\n")
        try:
            out = hinglish_selection_to_krutidev(raw)
        except Exception as e:
            QMessageBox.warning(self, "Convert failed", str(e))
            return
        fmt = self._kruti_char_format()
        c.beginEditBlock()
        c.insertText(out, fmt)
        c.endEditBlock()
        self._apply_kruti_display_after_convert()

    def _convert_document(self) -> None:
        plain = self.editor.toPlainText()
        try:
            out = hinglish_document_to_krutidev(plain)
        except Exception as e:
            QMessageBox.warning(self, "Convert failed", str(e))
            return
        fmt = self._kruti_char_format()
        self.editor.selectAll()
        c = self.editor.textCursor()
        c.beginEditBlock()
        c.insertText(out, fmt)
        c.endEditBlock()
        c.movePosition(QTextCursor.MoveOperation.Start)
        self.editor.setTextCursor(c)
        self._apply_kruti_display_after_convert()

    def _open_file(self) -> None:
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Open",
            "",
            "Word (*.docx);;HTML (*.html *.htm);;Text (*.txt);;All (*.*)",
        )
        if not path:
            return
        p = Path(path)
        if p.suffix.lower() == '.docx' and Document:
            try:
                doc = Document(p)
                text = '\\n'.join(para.text for para in doc.paragraphs)
                self.editor.setPlainText(text)
            except:
                QMessageBox.warning(self, 'Open failed', 'Could not read DOCX.')
                return
        else:
            text = p.read_text(encoding="utf-8", errors="replace")
            if p.suffix.lower() in (".html", ".htm"):
                self.editor.setHtml(text)
            else:
                self.editor.setPlainText(text)
        self._apply_document_page_geometry()

    def _save_file(self) -> None:
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Save",
            "",
            "Word (*.docx);;HTML (*.html);;Text (*.txt);;All (*.*)",
        )
        if not path:
            return
        p = Path(path)
        if p.suffix.lower() == '.docx' and Document:
            try:
                doc = Document()
                doc.add_paragraph(self.editor.toPlainText())
                doc.save(p)
            except:
                QMessageBox.warning(self, 'Save failed', 'Could not save DOCX.')
                return
        elif p.suffix.lower() == ".txt":
            p.write_text(self.editor.toPlainText(), encoding="utf-8")
        else:
            if not p.suffix:
                p = p.with_suffix(".html")
            p.write_text(self.editor.toHtml(), encoding="utf-8")


def main() -> int:
    app = QApplication(sys.argv)
    app.setApplicationName("HinglishKrutiEditor")
    app.setStyle("Fusion")
    apply_light_chrome(app)
    w = MainWindow()
    w.show()
    return app.exec()


if __name__ == "__main__":
    sys.exit(main())
